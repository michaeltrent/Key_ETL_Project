# -*- coding: utf-8 -*-
"""
Created on Thu Jun 11 08:11:18 2020

@author: mtrent

The purpose of this script is to process data from Keyland.
This will read in the data from an Excel SS, consolidate the data
into the pertinent fields, break the legal desritption into 
multiple lines on a section basis then search for the title
on the M drive and determine the nets on a section basis.

Steps:
    1. work on reading the data from Keyland export and consolidate where needed
    2. parse the legals and add lines where needed.
    3. find MOR for a given tract
    4. read the MOR and find the ownership for each owner
    5. update the SS and check that our net matches

"""

import pandas as pd
import numpy as np
import os
from os import listdir
from os.path import isfile, join
import re
import tkinter as tk
from tkinter import simpledialog
import docx
import math
from glob import glob
import win32com.client as win32
from win32com.client import constants
import shutil
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import os
import nltk
from nltk import word_tokenize
import cv2
from PyPDF2 import PdfFileWriter, PdfFileReader


#%% Convert the .doc files to .docx Code borrowed from StackOverflow

cleanDocs = False

if cleanDocs:
    # Create list of paths to .doc files
    path = 'C:/Users/micha/Documents/LTE/Keyland_Migration_Project'+'/'+ 'Clients (Team Folder)/' + Client +'/' + 'Title' 
    badPaths = glob(path + '/**\\*.doc', recursive=True)
    paths=[]
    
    for path in badPaths:
        paths.append(path.replace('\\', '/'))
    
    def save_as_docx(path):
        # Opening MS Word
        word = win32.gencache.EnsureDispatch('Word.Application')
        doc = word.Documents.Open(os.path.abspath(path))
        doc.Activate ()
    
        # Rename path with .docx
        new_file_abs = os.path.abspath(path)
        new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)
    
        # Save and Close
        word.ActiveDocument.SaveAs(
            new_file_abs, FileFormat=constants.wdFormatXMLDocument
        )
        doc.Close(False)
        try:
            shutil.move(path, path[0:path.rfind('/')]+'/Archived/'+path[path.rfind('/')+1::])
        except:
            print('Could not move file: ' + path[path.rfind('/')+1::])
    
    for path in paths:
            print('Converting MOR: ' + path[path.rfind('/')+1::])
            save_as_docx(path)


#%% Define the function to search for the MOR in the directory
''' This function will search for an MOR based on the tract legal description
and return the file path and name of the MOR to be opened:
    1. list contents of a given TR folder
    2. parse the legal description
    3. loop through the files in the folder to see if one contains the legal
    4. If the searchType is FULL then MOR must contain all tracts, otherwise it can contain partial tracts
    5. return the filepath and filename '''

def findMOR(Client, County, T, R, S, desc, searchType='FULL'):
    path = 'C:/Users/micha/Documents/LTE/Keyland_Migration_Project'+'/'+ 'Clients (Team Folder)/' + Client +'/' + 'Title'  + '\\' + County + '\\' + T + '-' + R
    fileList = [f for f in listdir(path) if isfile(join(path, f))]
    #Tokenize the description
    tracts = desc.split(',')
    #Determine the number of tracts in the legal
    numTracts = len(tracts)
    for file in fileList:
        #Initialize a counter for the number of tracts found 
        tractsFound = 0
        for tract in tracts:
            if file.count(tract.strip()) >= 1 and file.count('Sec ' + S) !=0 and file.count('MOR') == 1 and file.count('Chain') == 0 and file.lower().count('.docx') > 0:
                print('MOR found: ' + file)
                #If the MOR file name contains a legal tract update the tract found counter
                tractsFound += 1
                if tractsFound == numTracts and searchType == 'FULL': 
                    #If you have found a file with all tracts return the MOR
                    return path, file
                elif tractsFound >= 1:
                    return path, file
    failure = 'No MOR found with legal description', ''
    return failure
                
        
        
#%% Read data from MOR for each line

def readMOR(Client, County, T, R, S, desc, recNum, bkPg):
    ''' This function will call the findMOR function to determine
    which file to read, then read through on a lease basis and search for 
    the corresponding mineral interest and NMA values'''
    
    MI = 0
    NMA = 0
    path, MOR = findMOR(Client, County, T, R, S, desc, 'FULL')
    if path =='No MOR found with legal description':
        print('No full MOR match, searching partial tracts')
        path, MOR = findMOR(Client, County, T, R, S, desc, 'Part')
        if path =='No MOR found with legal description':
            print('WARNING, NO MOR FOUND')
            MI = np.nan
            NMA = np.nan
            MOR = path
            return MI, NMA, MOR
    doc = docx.Document(os.path.join(path, MOR))
    tableNum = np.nan
    rowNum = np.nan
    # We need run through the breakout tables and look for the recNum/BkPg
    # when we find the row that contains the recNum, we'll know the 
    # mineral ownership is in the second cell of that table/row.
    for tNum, table in enumerate(doc.tables):
        for rNum, row in enumerate(table.rows):
            for cNum, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    if para.text.find(recNum) > 0:
                        tableNum = tNum # This is the table
                        rowNum = rNum  # This is row
                        #print('Table ' + str(tableNum))
    
    #Pull the Mineral interst. If it's in percent format, remove the %
    #and convert it to a float, otherwise just convert it to float
    #Search each cell in the table/row to find the first paragraph that is a number
    #after removing any possible % characters by removing the last two characters
    #when you find that column (col) you have found your mineral interest
    #Then check to see if there are any additional paragraphs that would represent
    #revisions to the MOR by checking the next para or two for a float value
    miFound = False
    
    if not math.isnan(tableNum): 
        for cNum, cell in enumerate(doc.tables[tableNum].rows[rowNum].cells):
            if not miFound:
                for paraNum, para in enumerate(cell.paragraphs):
                    if len(para.text) > 0 and para.text[0:len(para.text)-2].replace('.', '',1).isdigit():
                        col = cNum
                        if len(cell.paragraphs) > paraNum:
                            for line in range(paraNum, len(cell.paragraphs)) :
                                if cell.paragraphs[line].text[0:len(cell.paragraphs[line].text)-2].replace('.','', 1).isdigit():
                                    MI = cell.paragraphs[line].text
                                    paraLine = line
                        else:
                            MI = cell.para.text
                            paraLine = paraNum
                        miFound = True
                        try:
                            if MI.find('%') > 0:
                                MI = float(MI[0:MI.find('%')])/100
                            else:
                                MI = float(MI)
                        except:
                            MI = 'Unable to locate MI for lease'
                        
                        try:                                    
                            NMA = float(doc.tables[tableNum].rows[rowNum].cells[col+1].paragraphs[paraLine].text)
                        except:
                            NMA = 'Unable to locate NMA for lease'
    if math.isnan(tableNum):
        MI = np.nan
        NMA = np.nan
        print('Warning, lease number ' + lease + ' not found in MOR')  

    return MI, NMA, MOR                  

#%% Define the read LPR function

''' This function will read in Keyland LPRs to extract owner phone numbers, 
contact info and check the interest from the MOR read. This will rely on the following
steps:
    1. Read in LPR'''

    

def LPROCR(path, saveLoc):
    
    #Set the path for tesseract
    
    pytesseract.pytesseract.tesseract_cmd = r'C:\Users\micha\anaconda3\tesseract\Library\bin\tesseract.exe'

    #Set the path for poppler

    poppler_path = 'C:\\Users\\micha\\anaconda3\\poppler\\Library\\bin'
    #Change the directory to save the file
    os.chdir(saveLoc)
    
    file = os.path.abspath(path)
    LPG = path[path.find('\\')+1:path.find('.pdf')]
    #Convert the pages to images
    pages = convert_from_path(file, 500, poppler_path = poppler_path)
    #Initialize a counter for the images
    imageCounter = 1
    
    #for each of the pages save the image as a jpeg
    for page in pages[0:2]:
        fileName = 'page_' + str(imageCounter) + '.jpeg'
        page.save(fileName, 'JPEG')
        imageCounter +=1
    
    filelimit = imageCounter-1
    #Create an empty text file
    outName = 'out_text'+LPG+'.txt'
    
    outfile = outName
    #For each of jpegs, convert the image to a string and append it to the file
    with open(os.path.join(os.path.abspath(saveLoc), outfile), 'a') as f:
        for i in range(1, filelimit+1):
            fileName = "page_"+str(i)+".jpeg"
            # Do some preprocessing to remove gridlines and noise
            img = cv2.imread(fileName)
            resultImage = img.copy()
            # Convert the image to grayscale
            grayimg = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            gray=cv2.cvtColor(img,cv2.COLOR_BGR2GRAY)
            linek = np.zeros((11,11),dtype=np.uint8)
            linek[...,5]=1
            #x=cv2.morphologyEx(gray, cv2.MORPH_OPEN, linek ,iterations=1)
            #gray-=x
            #cv2.imshow('gray',gray)
            #cv2.waitKey(0)
            img=gray
            #img = cv2.medianBlur(img,5)
            # Try to remove the grid lines to avoid breaking the pages into two
            laplacian = cv2.Laplacian(grayimg,cv2.CV_8UC1) # Laplacian Edge Detection
            minLineLength = 900
            maxLineGap = 100
            lines = cv2.HoughLinesP(laplacian,1,np.pi/180,100,minLineLength,maxLineGap)
            for line in lines:
                for x1,y1,x2,y2 in line:
                    cv2.line(grayimg,(x1,y1),(x2,y2),(255,255,255),1)
            ## Invert the image
            #ret,grayimg = cv2.threshold(grayimg,127,255,cv2.THRESH_BINARY_INV)
            ## Filter on threshold
            #img = cv2.threshold(grayimg, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
            ## Define the kernel
            kernel = np.ones((5,5),np.uint8)
            # Perform a morphological opening
            #img = cv2.morphologyEx(img, cv2.MORPH_OPEN, kernel)
            # Dilate the text
            #img = cv2.morphologyEx(img, cv2.MORPH_DILATE, kernel)
            text = str(((pytesseract.image_to_string(img)))) 
            #text = text.replace('-\n', '')
            f.write(text)
        # #Close the file
        # f.close()
    return

#%% Define a function to grab two pages of a final lease package and save them locally

def grabPages(inPath, outPath, numPages):
    #Grab the first two pages of each Final Package and save them locally
    pathSlice=inPath[inPath.find('Leases\\')+len('Leases\\')::]
    outPath = pathSlice[0:pathSlice.find('\\')+8] + '.pdf'
    outPath = outPath.replace('\\', '_')
    print('New filename: ' + outPath)
    finalPkg = PdfFileReader(open(inPath, "rb"))
    output = PdfFileWriter()
    for i in range(0, numPages):
        output.addPage(finalPkg.getPage(i))
        
    with open(outPath, "wb") as outputStream:
        output.write(outputStream)
        
#%% Define a stringNumSearch function to locate the beginning and a number in a string
''' This function will take a string input and output the beginning 
and end locations of a number with '-' delimeters. It will do this 
using a while loop and an interValue passed (neg or pos)'''
def stringNumSearch(string, delim, iterVal):
    # The tax ID could be in one of two formats: EIN XX-XXXXXX or SSN XXX-XX-XXXX. This search
    # should be able to handle both cases. First slice the text, then search
    # forward for the next instance of '-' which shuold be a delimeter in the 
    # tax ID. Then search backward from the delimeter checking each element
    # to determine if it is a digit. If you encounter two consecutive characters
    # that are not digits, then you have found the end of the ID. Add two to the 
    # index and we have the beginning. Similar process to find the end. This
    # should also work just fine for phone numbers.
    if string.count(delim) > 0 and string[string.find(delim)+1].isdigit() and string[string.find(delim)-1].isdigit():
       #for simplicity in code we'll define a vaiable with the Num loc
       numLoc = string.find(delim)
       #initialize a counter of non-digit elements
       notDigit = 0
       #initialize a counter for the while loop
       numSearch = iterVal
       while notDigit < 2:
           idStart = numLoc+numSearch
           if not string[idStart].isdigit():
               notDigit += 1
           numSearch += iterVal
       end = idStart
       notDigit = 0
       #initialize a counter for the while loop
       numSearch = iterVal
       while notDigit < 2:
           idStart = numLoc+numSearch
           if not string[idStart].isdigit():
               notDigit += 1
           numSearch += (-iterVal)
       beg = idStart + 1
       return string[beg:end]
    else:
       print('Error, the string passed does not caontain a valid number sequence')
       

#%% Snip of code to run the grab pages function

getPages =False

os.chdir('C:/Users/micha/Documents/LTE/Keyland_Migration_Project/LPRs')
if getPages:
    for path in paths:
        print('Processing: ' + path)
        
        grabPages(path, 'C:/Users/micha/Documents/LTE/Keyland_Migration_Project/LPRs', 2)
    
#%% Run the OCR

runOCR = True

if runOCR:
    
    path = 'C:/Users/micha/Documents/LTE/Keyland_Migration_Project/LPRs' 
    paths =[f for f in glob(path + '/**\\*.pdf', recursive=True)]
    
    for path in paths[186::]:
        print('Processing: ' + path)
        LPROCR(path, 'C:/Users/micha/Documents/LTE/Keyland_Migration_Project/LPRs/LPG_Texts')

#%% If needed re-name the files once OCRed to include RecNum

LPRTextRename = True

if LPRTextRename:

    #Find the LPGs on the M Drive:
    path = 'C:/Users/micha/Documents/LTE/Keyland_Migration_Project/LPRs' 
    
    #paths =[f for f in glob(path + '/**\\*.pdf', recursive=True) if 'Final Lease Package Sent' in f]
    paths =[f for f in glob(path + '/**\\*.txt', recursive=True)]
    
    for path in paths:
        with open(path, 'r', errors='ignore') as file:
            text = file.read()#.replace('\n', '')
            #Define a parameter to find the first string of 6 digits after the Bk-Pg
            #character. That will be the rec num.
            recLoc = text.find('Bk-Pg')
            if recLoc == -1:
                recLoc = text.upper().find('RECORDING INFO')
            if recLoc !=-1:
                recFound = False
                numIter = 0
                while not recFound:
                    if text[recLoc:recLoc+6].isdigit():
                        recNum = text[recLoc:recLoc+6]
                        recFound = True
                    numIter += 1
                    recLoc += 1
                    if numIter == 100:
                        #if the num iters is high, then break out of the while loop. 
                        recNum = np.nan
                        break
        try:
            recNum = int(recNum)
            if not math.isnan(int(recNum)):
                print('Renaming file: ' + path)
                print('New file name: ' + path[0:path.find('.txt')]+str(recNum)+'.txt')
            
                os.rename(path, os.path.abspath(path[0:path.find('.txt')]+str(recNum)+'.txt'))
            else:
                print('No reception number found for: ' + path)
        except:
            print('No Rec Number found for file: ' + path)
        
        try:
            basePath = path[0:path.rfind('\\')]
            fileName = path[path.rfind('\\')+1::]
            #shutil.move(path, os.path.abspath(basePath + '/' + 'No_Rec_Num/' + fileName))
        except:
            print('Counld Not Move: ' + path)
            

        
#%%Function to read the output from the OCRed LPRs 

# #Find the LPGs on the M Drive:
# path = 'C:/Users/micha/Documents/LTE/Keyland_Migration_Project/LPRs' 

# #paths =[f for f in glob(path + '/**\\*.pdf', recursive=True) if 'Final Lease Package Sent' in f]
# paths =[f for f in glob(path + '/**\\*.txt', recursive=True)]

# with open(paths[1], 'r', errors='ignore') as file:
#     text = file.read()#.replace('\n', '')
#     # Create a slice of the text to search for Tax ID number. 
#     taxSlice = text[text.find('Tax')::]
#     taxID = stringNumSearch(taxSlice, '-', 1)
#     # If the first '-' is not bounded by digits then we have an exception
#     phoneSlice = text[text.find('Phone')::]
#     phoneNum = stringNumSearch(phoneSlice, '-', 1)

#%% Load and clean data

#Eventually need to insert dialog boxes for excelName and dataPath
#for now, I'll hard code it in. 

Client = 'KODA'
County = 'Williams'

dataPath = 'C:/Users/micha/Documents/LTE/Keyland_Migration_Project/'
excelName = 'KODA_Union_Leases.xlsx'

#Load the data and remove the carriage returns
klData = pd.read_excel(os.path.join(dataPath, excelName))

klData = klData.replace(r'\n', ' ', regex = True)

#Add Columns for T, R, S, Desc, and MOR

klData.insert(2, 'Township', np.nan)
klData.insert(3, 'Range', np.nan)
klData.insert(4, 'Sec', np.nan)
klData.insert(5, 'Legal', np.nan)
klData.insert(6, 'MOR', np.nan)
klData.insert(7, 'Phone Number', np.nan)
klData.insert(8, 'Tax ID', np.nan)
# Insert a column for mineral interest and NMA from MOR

klData.insert(klData.columns.get_loc('NetAcresAdding')+1, 'MineralInterest', np.nan)
klData.insert(klData.columns.get_loc('MineralInterest')+1, 'MORCalcNMA', np.nan)
del klData['MarkProblemGraphic']
del klData['Description_Lease']

data = pd.DataFrame(columns = klData.columns)

#%% Parse the legal to create a full legal for each line

LPRpath = 'C:/Users/micha/Documents/LTE/Keyland_Migration_Project/LPRs' 
LPRpaths =[f for f in glob(LPRpath + '/**\\*.txt', recursive=True)]

for row in range(0,klData.shape[0]):
    legal = klData.iloc[row, klData.columns.get_loc('Description_Lease_2')]
    print(legal)
    #TRS will be the data preceeding the first instance of Sec.
    TRS = legal[0:legal.find('Sec')-1]
    #Township and Range split
    T = TRS.split('-')[0]; R = TRS.split('-')[1]
    #The legal will be everything commencing with first instance of 'Sec'
    legal = legal[legal.find('Sec')::]
    #We need the county on a lease basis
    County = klData.iloc[row, klData.columns.get_loc('PROSPECT::prospect county calc')].split(' ')[0]
    for Section in range(0, legal.count('Sec')):
        #For each section, find the legal for the first section
        #If there is more than one section in legal if not just grab
        #the section legal
        if legal.count('Sec') > 1:
            #If there is more than one section, then we only want the string from the first section to the beginning of the next
            Sec = legal[0:legal[1::].find('Sec')]
            #Trim the legal of the first sec for processing subsequent sections recursively
            legal = legal[legal.find('Sec')+1::]
            #print('more than one Sec: ' + Sec)
        else:
            #If there is only one section, then we don't need to trim the string
            Sec = legal
            #print('one section: ' + Sec)
        S = Sec[3:Sec.find(":")]
        S = S.lstrip()
        if len(S) < 2:
            S = '0'+S
        legal = legal[legal.find('Sec')::]
        desc = Sec[Sec.find(':')+2::]
        phoneNum = np.nan
        taxID = np.nan
        #Copy the data from the KL Data set to append into the new data set
        newRow = klData.loc[klData['LeaseNumberCalc'] == klData.iloc[row, klData.columns.get_loc('LeaseNumberCalc')]]
        #If there is no Rec Number, then replace the NaN with 999999
        if not math.isnan(klData.iloc[row, klData.columns.get_loc('RecordingInfo')]):
            lease = str(int(klData.iloc[row, klData.columns.get_loc('RecordingInfo')]))
        else:
            lease = '999999'
        try:
            MI, NMA, MOR = readMOR(Client, County, T, R, S, desc, lease, 'BkPg')
        except:
            MI = 'No MOR Found'
            NMA = 'No MOR Found'
            MOR = 'No MOR Found'
            print('Unable to process lease ' + lease)
        try:
            for path in LPRpaths:
                if path.count(str(int(lease))):
                    with open(path, 'r', errors='ignore') as file:
                        text = file.read()
                        # Create a slice of the text to search for Tax ID number. 
                        taxSlice = text[text.find('Tax')::]
                        taxID = stringNumSearch(taxSlice, '-', 1)
                        # If the first '-' is not bounded by digits then we have an exception
                        phoneSlice = text[text.find('Phone')::]
                        phoneNum = stringNumSearch(phoneSlice, '-', 1)
        except:
            phoneNum = np.nan
            taxID = np.nan
        newRow['Township'] = T
        newRow['Range'] = R
        newRow['Sec'] = S
        newRow['Legal'] = desc
        newRow.iloc[0,newRow.columns.get_loc('MineralInterest')] = MI
        newRow['MORCalcNMA'] = NMA
        newRow['MOR'] = MOR
        newRow['Phone Number'] = phoneNum
        newRow['Tax ID'] = taxID
        print(newRow['Phone Number'])
        #print('New Leagal is: '+T+R+S+desc)
        data = data.append(newRow, ignore_index = True)

#%% Save data in excel
#del data['Description_Lease']
os.chdir('C:/Users/micha/Documents/LTE/Keyland_Migration_Project')
#del data['Description_Lease_2']

newExcelName = 'KODA_UNION_LEASES_NEW_FORMAT.xlsx'
data.to_excel(os.path.join(dataPath, newExcelName), index = False)

    